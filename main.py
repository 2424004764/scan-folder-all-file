import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import os
import threading
from consts import *
import sys
import fitz
from win32com import client as wc
import docx
from docx import Document


def delete_item(item_id, tree):
    # 根据 item_id 删除相应的行
    tree.delete(item_id)


class ScanFolderFiles:
    # 是否继续扫描
    continue_scan = True
    # 文件包含的敏感词映射
    file_sensitive_words_map = {
        # 示例结构
        # file_path: [敏感词列表]
    }

    def create_treeview(self, _root, _w):
        """
        创建一个tree_view组件
        :param _root: 需要放置到哪个frame下
        :param _w: 使用哪个表头
        :return:
        """
        header = TREEVIEW_TABLE_HEADER[_w]
        _tree_var = ttk.Treeview(_root, columns=tuple(item["key"] for item in header), show='headings')
        for item in header:
            _tree_var.column(item["key"], width=item["width"], minwidth=item["min_width"], anchor=item["anchor"])
            _tree_var.heading(item["key"], text=item["text"])

        _tree_var.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # 绑定鼠标单击事件
        _tree_var.bind("<ButtonRelease-1>",
                       lambda event: self.on_tree_select(event, _tree_var))
        return _tree_var

    def set_ico(self):
        """
        设置ico
        :return:
        """
        # 获取打包后可执行文件所在目录的绝对路径
        if getattr(sys, 'frozen', False):
            # 打包后的可执行文件路径
            base_path = sys._MEIPASS
        else:
            # 普通 Python 程序执行的路径
            base_path = os.path.abspath(".")
        # 图标文件相对于可执行文件的路径
        icon_path = os.path.join(base_path, LOGO_ICO_PATH)
        # 设置窗口图标
        if os.path.exists(icon_path):
            self.root.iconbitmap(icon_path)

    # 定义样式
    def init_style(self):
        # 定义一个新的样式，名为 "RedBorderLabel.TLabel"，继承自 "TLabel"
        self.red_style = ttk.Style()
        self.red_style.configure("RedBorderLabel.TLabel", foreground="red")

    def on_tree_select(self, event, tree):
        item_id = tree.focus()  # 获取当前焦点行的ID
        if item_id:  # 确保有选中行
            col_index = tree.identify_column(event.x)  # 获取点击列的索引
            # col_heading = tree.heading(col_index)['text']  # 获取点击列的标题文本
            values = tree.item(item_id, 'values')  # 获取选中行的值

            is_continue = False
            if tree == self.file_listbox:
                # 如果点击的是已扫描文件列表
                if col_index == "#3":
                    # 忽略操作
                    delete_item(item_id, tree)
                    is_continue = True
            if tree == self.error_listbox:
                operable_index = ["#4", "#5"]
                if not col_index in operable_index:
                    # 点击的不是删除或忽略
                    return
                # 如果点击的是异常文件列表
                if col_index == operable_index[0]:
                    # 忽略操作
                    delete_item(item_id, tree)
                    is_continue = True
                if col_index == operable_index[1]:
                    # 删除操作
                    confirmed = messagebox.askyesno("提示", "是否将敏感词从文件内容中删除？")
                    if not confirmed:
                        return
                    try:
                        self.delete_file_content_by_words(values[1])
                        delete_item(item_id, tree)
                        is_continue = True
                    except OSError as e:
                        messagebox.showinfo("删除失败", f"原因：{e}")

            if not is_continue:
                # 点击的不是删除或者忽略
                return
            # 判断是删除的已扫描文件还是异常文件
            if tree == self.file_listbox:
                self.file_count -= 1
                self.update_file_count()  # 更新已扫描文件数
            if tree == self.error_listbox:
                self.error_count -= 1
                self.update_error_count()  # 更新异常文件数
            self.update_tree_view_index(tree)  # 重排索引

    def __init__(self, root):
        self.error_listbox = None
        self.red_style = None
        self.file_count = 0
        self.folder_count = 0
        self.error_count = 0
        self.root = root
        self.root.title(SOFT_NAME)

        self.set_ico()
        self.init_style()

        mainframe = ttk.Frame(self.root, padding="3 3 12 12")
        mainframe.pack(fill=tk.BOTH, expand=True)

        # 第一行：文件夹路径输入框和选择按钮
        folder_frame = ttk.Frame(mainframe)
        folder_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(folder_frame, text='文件夹:').pack(side=tk.LEFT)
        self.folder_path = tk.StringVar()
        ttk.Entry(folder_frame, state="readonly", width=70, textvariable=self.folder_path).pack(side=tk.LEFT, padx=5)
        self.select_folder_btn = ttk.Button(folder_frame, text="选择文件夹", command=self.select_folder)
        self.select_folder_btn.pack(side=tk.LEFT, padx=5)

        # 第二行：文件列表框和滚动条
        listbox_frame = ttk.Frame(mainframe)
        listbox_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        ttk.Label(listbox_frame, text="已扫描文件名列表:").pack(side=tk.TOP, anchor=tk.W)
        self.file_listbox = self.create_treeview(listbox_frame, TABLE_HEADER_SUCCESS)

        scrollbar = ttk.Scrollbar(listbox_frame, orient=tk.VERTICAL, command=self.file_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.file_listbox.config(yscrollcommand=scrollbar.set)

        # 第三行：显示已扫描的文件夹数和文件数的标签
        status_frame = ttk.Frame(mainframe)
        status_frame.pack(fill=tk.X, padx=5, pady=5)

        # 正在扫描的文件label
        self.current_scan_file_label = ttk.Label(status_frame, text="")
        self.current_scan_file_label.pack(side=tk.TOP, anchor=tk.W)

        self.folder_count_label = ttk.Label(status_frame, text="已扫描文件夹数: 0")
        self.folder_count_label.pack(side=tk.LEFT)

        self.file_count_label = ttk.Label(status_frame, text="已扫描文件数: 0")
        self.file_count_label.pack(side=tk.LEFT, padx=10)

        self.export_button = ttk.Button(status_frame, text="导出已扫描文件", command=self.export_normal_to_txt)
        self.export_button.pack(side=tk.LEFT, padx=10)

        # 第五行：异常文件名记录文本框、滚动条和异常文件数标签
        error_frame = ttk.Frame(mainframe)
        error_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        ttk.Label(error_frame, text="异常文件名列表:").pack(side=tk.TOP, anchor=tk.W)
        self.error_listbox = self.create_treeview(error_frame, TABLE_HEADER_FAIL)

        self.error_scrollbar = ttk.Scrollbar(error_frame, orient=tk.VERTICAL, command=self.error_listbox.yview)
        self.error_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.error_listbox.config(yscrollcommand=self.error_scrollbar.set)

        self.error_count_label = ttk.Label(mainframe, text="异常文件数: 0")
        self.error_count_label.pack(side=tk.LEFT, pady=5)

        self.export_error_button = ttk.Button(mainframe, text="导出异常文件", command=self.export_error_to_txt)
        self.export_error_button.pack(side=tk.LEFT, padx=10)

        self.stop_button = ttk.Button(mainframe, text="停止扫描", command=self.stop_scanning)
        self.stop_button.pack(side=tk.LEFT, padx=10)
        self.stop_button.config(state='disabled')  # 默认不可点击，扫描时可点击。扫描完成后不可点击

        # 初始化已扫描完成的标签但不显示
        self.scan_completed_label = ttk.Label(mainframe, text="", style="RedBorderLabel.TLabel")
        self.scan_completed_label.pack(pady=5)

        # 暂时先禁用导出按钮，扫描完成再启用
        self.disable_export_button()

    def enable_export_button(self):
        """
        启用导出按钮，扫描完成后才能点击
        :return:
        """
        self.export_button.config(state='normal')
        self.export_error_button.config(state='normal')

    def disable_export_button(self):
        """
        禁用导出按钮，扫描完成后才能点击
        :return:
        """
        self.export_button.config(state='disabled')  # 默认不可点击，扫描完成可点击
        self.export_error_button.config(state='disabled')  # 默认不可点击，扫描完成可点击

    def update_tree_view_index(self, _w):
        """
        更新tree_view组件索引，因为删除元素了如果不重排索引就会出现 2、3、4这种数据
        :param _w: tree_view组件实例
        :return:
        """
        children = _w.get_children()
        file_len = len(children)
        if file_len == 0:
            # 说明被删完了
            return
        for i in range(file_len):
            raw_values = _w.item(children[i])['values']  # 这一行原值，是一个数组
            raw_values[0] = i + 1
            _w.item(children[i], values=raw_values)

    def stop_scanning(self):
        self.continue_scan = False
        self.stop_button.pack_forget()  # 按钮默认不可见，只有开始扫描才可见

    def delete_item(self, item_id):
        print(item_id)

    def select_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.continue_scan = True
            # 选择了文件夹后 将按钮置为不可用，等全部读取后再置为可用
            self.select_folder_btn.config(state='disabled')
            # 将停止扫描按钮置为可见
            self.stop_button.pack(side=tk.LEFT, padx=10)
            # 设置选择的文件夹
            self.folder_path.set(folder_path)
            # 将完成提示清空
            self.hide_scan_completed()
            self.stop_button.config(state='normal')
            # 暂时先禁用导出按钮，扫描完成再启用
            self.disable_export_button()
            threading.Thread(target=self.process_folder, args=(folder_path,)).start()

    def process_folder(self, folder_path):
        # 清空已扫描文件列表
        self.file_listbox.delete(*self.file_listbox.get_children())
        self.error_listbox.delete(*self.error_listbox.get_children())
        self.folder_count = 0
        self.file_count = 0
        self.error_count = 0

        scan_count = 1
        scan_err_count = 1
        # 遍历文件夹中的所有文件
        for root, dirs, files in os.walk(folder_path):
            if not self.continue_scan:
                break
            self.folder_count += 1
            self.update_folder_count()  # 更新已扫描文件夹数

            for file_name in files:
                if not self.continue_scan:
                    break
                file_path = os.path.join(root, file_name)
                # 更新界面显示文件名
                self.update_file_list(scan_count, file_path)
                scan_count += 1
                self.file_count += 1
                self.update_file_count()  # 更新已扫描文件数

                # 检查文件内容是否包含敏感词
                if self.exist_sensitive_words(file_path):
                    self.update_error_text(scan_err_count, file_path)
                    scan_err_count += 1
                    self.error_count += 1
                    self.update_error_count()  # 更新异常文件数

        # 更新扫描完成标签
        self.show_scan_completed()
        # 恢复选择文件按钮为正常状态
        self.select_folder_btn.config(state='normal')
        # 将停止扫描按钮置为不可点击状态，因为扫描完成了
        self.stop_button.config(state='disabled')
        # 启用导出按钮
        self.enable_export_button()
        # 更新正在扫描的文件为空
        self.current_scan_file_label.config(text="")

    def delete_file_content_by_words(self, file_path):
        """
        将某个异常文件中的敏感词全部删除
        :param file_path:
        :return:
        """
        file_extension = self.get_file_extension(file_path)
        if not file_extension in CHECK_FILE_SUFFIX:
            # 不在检查范围内的文件类型直接忽略
            return
        if file_extension == 'txt':
            self.delete_file_content_by_txt(file_path)
        elif file_extension == 'pdf':
            self.delete_file_content_by_pdf(file_path)
        elif file_extension == 'doc':
            # 先将doc另存为docx，再走docx一样的处理流程
            tem_docx_path = self.doc_save_as_docx(file_path)
            self.delete_file_content_by_docx(file_path, tem_docx_path)
            try:
                # 删除临时文件
                os.unlink(tem_docx_path)
            except FileNotFoundError as e:
                print('删除临时文件报错：', e)
                pass
        elif file_extension == 'docx':
            self.delete_file_content_by_docx(file_path, file_path)

    def delete_file_content_by_docx(self, raw_file_path, tem_file_path):
        # 打开 Word 文档
        doc = Document(tem_file_path)

        # 遍历文档中的每个段落
        for paragraph in doc.paragraphs:
            for keyword in self.file_sensitive_words_map[raw_file_path]:
                if keyword in paragraph.text:
                    # 替换关键字
                    updated_text = paragraph.text.replace(keyword, "")
                    paragraph.text = updated_text

        # 覆盖源文档
        doc.save(raw_file_path)

    def delete_file_content_by_txt(self, file_path):
        """
        从txt源文件中删除所有关键字
        :param file_path:
        :return:
        """
        # 打开文件并读取内容
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()

        # 遍历关键字列表，将每个关键字替换为空字符串
        for keyword in self.file_sensitive_words_map[file_path]:
            content = content.replace(keyword, '')

        # 将更新后的内容写回到文件中
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)

    def delete_file_content_by_pdf(self, file_path):
        # 打开 PDF 文件
        pdf_document = fitz.open(file_path)

        # 遍历每一页
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]

            # 遍历关键字列表，查找并删除关键字
            for keyword in self.file_sensitive_words_map[file_path]:
                text_instances = page.search_for(keyword)
                for inst in text_instances:
                    # 将关键字替换为空字符串
                    page.add_redact_annot(inst)
                    page.apply_redactions()

        # 保存修改后的 PDF 文件到临时文件
        file_name = self.get_file_name_without_extension(file_path)
        new_file_path = os.path.dirname(file_path)
        # 给临时文件命名
        file_new_name = "{}\\__{}.pdf".format(new_file_path, file_name)
        pdf_document.save(file_new_name)
        pdf_document.close()

        # 将临时文件重命名为源文件
        os.replace(file_new_name, file_path)

    def get_file_name_without_extension(self, file_path):
        # 使用 os.path.splitext 分割文件名和扩展名
        file_name = os.path.basename(file_path)
        file_name_without_extension = os.path.splitext(file_name)[0]
        return file_name_without_extension

    def doc_save_as_docx(self, file_path):
        word = wc.Dispatch('Word.Application')
        # 替换，否则会因为有空格而读取失败
        file_path = file_path.replace('/', '\\')
        doc = word.Documents.Open(file_path)  # 目标路径下的文件
        file_name = self.get_file_name_without_extension(file_path)
        new_file_path = os.path.dirname(file_path)
        # 给临时文件命名
        file_new_name = "{}\\__{}.docx".format(new_file_path, file_name)
        doc.SaveAs(file_new_name, 12, False, "", True, "", False, False, False,
                   False)  # 转化后路径下的文件
        doc.Close()
        word.Quit()
        # 临时文件，扫描完成后删除
        return file_new_name

    def exist_sensitive_words(self, file_path):
        file_extension = self.get_file_extension(file_path)
        if not file_extension in CHECK_FILE_SUFFIX:
            # 不在检查范围内的文件类型直接忽略
            return False
        if file_extension == 'txt':
            return self.check_by_txt(file_path)
        elif file_extension == 'pdf':
            return self.check_by_pdf(file_path)
        elif file_extension == 'doc':
            # 先将doc另存为docx，再走docx一样的处理流程
            # 因为doc不能直接打开，所以先转成docx
            tem_docx_path = self.doc_save_as_docx(file_path)
            flag = self.check_by_docx(file_path, tem_docx_path)
            try:
                # 删除临时文件
                os.unlink(tem_docx_path)
            except FileNotFoundError as e:
                print('删除临时文件报错：', e)
                pass
            return flag
        elif file_extension == 'docx':
            return self.check_by_docx(file_path, file_path)
        return False

    def check_by_docx(self, real_file_path, tem_file_path):
        file = docx.Document(tem_file_path)
        # 输出每一段的内容
        flag = False
        for para in file.paragraphs:
            if len(para.text) == 0:
                continue
            tem_flag = self.map_sensitive_words(real_file_path, para.text)
            if not flag:
                # 防止最后一页没有敏感词从而将总的结果也判定为无敏感词
                flag = tem_flag
        return flag

    def check_by_pdf(self, file_path):
        try:
            # 打开 PDF 文件
            pdf_document = fitz.open(file_path)
        except fitz.FileDataError as e:
            print("{} 打开失败，可能文件已损坏，尝试用wps能否打开，错误原因：{}".format(file_path, e))
            return False

        # 读取每一页的文本内容
        all_text = ""
        for page_number in range(len(pdf_document)):
            page = pdf_document[page_number]
            text = page.get_text()
            all_text += text

        # 关闭 PDF 文件
        pdf_document.close()
        if all_text:
            return self.map_sensitive_words(file_path, all_text)
        return False

    def check_by_txt(self, file_path):
        try:
            flag = False
            with open(file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    # 处理每一行内容
                    content = line.strip()
                    if len(content) == 0:
                        continue
                    # 检查关键字是否在文件内容中
                    tem_flag = self.map_sensitive_words(file_path, content)
                    if not flag:
                        # 防止最后一页没有敏感词从而将总的结果也判定为无敏感词
                        flag = tem_flag
            return flag
        except FileNotFoundError:
            print(f"File '{file_path}' not found.")
            return False

    def map_sensitive_words(self, file_path, content):
        flag = False
        for keyword in ERROR_WORDS:
            if keyword in content:
                if file_path in self.file_sensitive_words_map:
                    if keyword not in self.file_sensitive_words_map[file_path]:
                        self.file_sensitive_words_map[file_path].append(keyword)
                else:
                    self.file_sensitive_words_map[file_path] = [keyword]
                flag = True
        return flag

    def get_file_extension(self, file_path):
        # 使用 os.path.splitext() 获取文件名和后缀名
        _, file_extension = os.path.splitext(file_path)
        # 去除后缀名中的点号，并转换为小写
        return file_extension.lstrip('.').lower()

    def update_file_list(self, scan_count, file_path):
        self.file_listbox.insert('', 'end', values=(scan_count, file_path, '忽略'))
        self.file_listbox.yview_moveto(1.0)  # 滚动到最底部
        # 更新正在扫描的文件
        self.current_scan_file_label.config(text=f"正在扫描文件: {file_path}")

    def update_folder_count(self):
        self.folder_count_label.config(text=f"已扫描文件夹数: {self.folder_count}")

    def update_file_count(self):
        self.file_count_label.config(text=f"已扫描文件数: {self.file_count}")

    def update_error_text(self, scan_err_count, file_path):
        words = ','.join(self.file_sensitive_words_map[file_path])
        self.error_listbox.insert('', 'end', values=(scan_err_count, file_path, words, '忽略', '删除'))
        self.error_listbox.yview_moveto(1.0)  # 滚动到最底部

    def update_error_count(self):
        self.error_count_label.config(text=f"异常文件数: {self.error_count}")

    def show_scan_completed(self):
        self.scan_completed_label.config(text="已扫描完成!")

    def hide_scan_completed(self):
        self.scan_completed_label.config(text="")

    def export_normal_to_txt(self):
        file_children = self.file_listbox.get_children()
        file_len = len(file_children)
        if file_len == 0:
            messagebox.showinfo("提示", f"已扫描的文件列表为空")
            return
        output_file = filedialog.asksaveasfilename(defaultextension=".txt", initialfile="success_file_names",
                                                   filetypes=[("Text Files", "*.txt")])
        if output_file:
            with open(output_file, "w", encoding='utf-8') as f:
                for i in range(file_len):
                    file_path = self.file_listbox.item(file_children[i]).get('values')[1]
                    f.write(file_path + "\n")
            messagebox.showinfo("导出成功", f"已将正常文件名导出到：{output_file}")

    def export_error_to_txt(self):
        err_children = self.error_listbox.get_children()
        err_len = len(err_children)
        if err_len == 0:
            messagebox.showinfo("提示", f"扫描的异常文件列表为空")
            return
        output_file = filedialog.asksaveasfilename(defaultextension=".txt", initialfile="error_file_names",
                                                   filetypes=[("Text Files", "*.txt")])
        if output_file:
            with open(output_file, "w", encoding='utf-8') as f:
                for i in range(err_len):
                    file_path = self.error_listbox.item(err_children[i]).get('values')[1]
                    words = ','.join(self.file_sensitive_words_map[file_path])
                    f.write(f"文件：{file_path} 包含敏感词：{words}\n")
            messagebox.showinfo("导出成功", f"已将异常文件名导出到：{output_file}")


def main():
    root = tk.Tk()
    ScanFolderFiles(root)
    root.mainloop()


if __name__ == "__main__":
    main()
